import streamlit as st
import pandas as pd
import os
import json
import re
import plotly.express as px
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import multiprocessing

# --- WINDOWS MULTIPROCESSING PROTECTION ---
if __name__ == '__main__':
    multiprocessing.freeze_support()

# --- SELF-HEALING: Folder Initialization ---
for path in ["data/raw", "data/working", "data/output", "data/config"]:
    if not os.path.exists(path):
        os.makedirs(path)

# --- CORE PIPELINE IMPORTS ---
try:
    from src.ingest import normalize_csv_parallel
    from src.tagger import tag_messages
    from src.incidents import build_incident_report
except ImportError as e:
    st.error(f"❌ Core Logic Missing: {e}. Ensure 'src' folder exists.")

# --- EXPORT ENGINE: TWO-PASS WORD GENERATION ---
def export_to_word(incident_csv, output_path):
    if not os.path.exists(incident_csv):
        return False
    df = pd.read_csv(incident_csv)
    df['dt'] = pd.to_datetime(df['dt'])
    df = df.sort_values(by=['category', 'dt'])
    doc = Document()

    # Pass 1: Header
    title = doc.add_heading('CERTIFIED REPORT OF CUSTODY INCIDENTS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Pass 1: Exhibit Index Summary
    doc.add_heading('Exhibit Index Summary', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text = 'Exhibit ID', 'Date', 'Legal Category'
    for cell in hdr_cells:
        for p in cell.paragraphs: p.runs[0].bold = True
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text, row_cells[1].text = str(row['exhibit_id']), row['dt'].strftime('%Y-%m-%d %H:%M')
        row_cells[2].text = str(row['category']).upper().replace('_', ' ')

    # Pass 2: Categorized Chapters
    for category in df['category'].unique():
        doc.add_page_break()
        doc.add_heading(f"CHAPTER: {category.replace('_', ' ').upper()}", level=1)
        for _, row in df[df['category'] == category].iterrows():
            p = doc.add_paragraph()
            p.add_run(f"EXHIBIT {str(row['exhibit_id']).upper()}").bold = True
            
            meta = doc.add_paragraph()
            meta_run = meta.add_run(f"VERIFIED SOURCE | Row: {row.get('raw_row_number')} | ID: {row.get('message_id')}")
            meta_run.font.size, meta_run.italic = Pt(8), True
            
            doc.add_paragraph("Evidence Quote:").runs[0].bold = True
            q = doc.add_paragraph()
            q.add_run(f"\"{row.get('evidence_quote', row.get('text', ''))}\"").italic = True
            q.paragraph_format.left_indent = Inches(0.75)
            
            reason = doc.add_paragraph()
            reason.add_run(f"Legal Reasoning: {row['reasoning']}").bold = True
            doc.add_paragraph("_" * 60)
    doc.save(output_path)
    return True

# --- UI BRANDING & DEFAULTS ---
LEGAL_NAVY = "#002D62"; LEGAL_GOLD = "#D4AF37"; BG_LIGHT = "#F8F9FA"
COURT_CATEGORIES = ["Parenting Time Interference", "Refusal to Cooperate", "Radio Silence / Communication Gaps", "Medical / Well-being Issues", "Parental Alienation Behavior", "Education / School Issues", "Hostile / Harassing Tone", "Feigned Ignorance"]
DEFAULT_IDS = "Joe, Giuseppe, Dad, Me, Myself"
DEFAULT_KEYWORDS = ["pickup", "dropoff", "doctor", "dentist", "school", "grades", "sick", "late", "court", "order", "refuse"]

st.set_page_config(page_title="STRICT v3.8.5 | Custody Agent", layout="wide")
st.markdown(f"<style>.main {{ background-color: {BG_LIGHT}; }} .stButton>button {{ background-color: {LEGAL_NAVY}; color: white; border-radius: 5px; font-weight: bold; width: 100%; }} div[data-testid='stExpander'] {{ background-color: white; border-left: 5px solid {LEGAL_NAVY}; }}</style>", unsafe_allow_html=True)

# --- SIDEBAR ---
with st.sidebar:
    st.title("⚖️ STRICT v3.8.5")
    uploaded_file = st.file_uploader("1. Upload Text Export (CSV)", type="csv")
    if uploaded_file and st.button("🔍 KICK OFF PRELIMINARY SCAN"):
        raw_path = "data/raw/imazing_export.csv"
        with open(raw_path, "wb") as f: f.write(uploaded_file.getbuffer())
        df_peek = pd.read_csv(raw_path, nrows=1000)
        st.session_state['suggestions'] = {"senders": df_peek['Sender Name'].dropna().unique().tolist()}
        st.rerun()

    st.divider()
    st.header("⚙️ Analysis Tuning")
    gap_threshold = st.slider("Radio Silence Threshold (Hours)", 24, 168, 72)
    buffer_hours = st.number_input("Exchange Buffer (Hours)", 1, 12, 4)
    flag_intensity = st.select_slider("AI Flagging Sensitivity", options=["Conservative", "Balanced", "Aggressive"], value="Balanced")

# --- MAIN DASHBOARD ---
if uploaded_file:
    tab1, tab2, tab3, tab4 = st.tabs(["🚀 COMMAND CENTER", "⚖️ LEGAL SCOPE", "🔍 AUDIT STATION", "📈 ANALYTICS"])

    with tab2:
        st.subheader("Define Analysis Scope")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("#### ⚖️ Court-Compliant Categories")
            selected_cats = [cat for cat in COURT_CATEGORIES if st.checkbox(cat, value=True)]
            custom_cat = st.text_input("Add Custom Category:", placeholder="e.g. Financial Non-Compliance")
            if custom_cat: selected_cats.append(custom_cat)
        with col2:
            st.markdown("#### 🔑 Conflict Keywords")
            selected_kws = [kw for kw in DEFAULT_KEYWORDS if st.checkbox(kw, value=True, key=f"kw_{kw}")]
            st.divider()
            custom_kw_input = st.text_area("✍️ Custom Keywords to Hunt:", placeholder="vacation, passport, behavior")
            if custom_kw_input: selected_kws.extend([k.strip() for k in custom_kw_input.split(",")])

    with tab1:
        st.subheader("Identity & Schedule History")
        c1, c2 = st.columns(2)
        with c1: me_names = st.text_area("Names for YOU:", value=st.session_state.get('me_names_list', DEFAULT_IDS))
        with c2: you_names = st.text_area("Names for THEM:", value="Mom, Mother")
        
        st.divider()
        st.subheader("📅 Varied Schedule Phases")
        if 'phases' not in st.session_state:
            st.session_state['phases'] = [{"Phase": "Initial Agreement", "Start": "2024-01-01", "End": "", "Style": "Standard"}]
        phase_df = st.data_editor(pd.DataFrame(st.session_state['phases']), num_rows="dynamic", use_container_width=True)

        if st.button("🚀 EXECUTE FULL ANALYSIS"):
            with st.status("Analyzing...", expanded=True):
                normalize_csv_parallel("data/raw/imazing_export.csv", "data/working/messages_normalized.csv", [n.strip() for n in me_names.split(",")], [n.strip() for n in you_names.split(",")])
                tag_messages(pd.read_csv("data/working/messages_normalized.csv"), custom_keywords=selected_kws, selected_categories=selected_cats).to_csv("data/output/messages_tagged.csv", index=False)
                build_incident_report("data/output/messages_tagged.csv", "data/working/messages_normalized.csv", gap_threshold_hours=gap_threshold)
            st.success("Analysis Complete. Check Audit Station.")

    with tab3:
        st.subheader("Evidence Audit Station")
        if os.path.exists("data/output/incident_index.csv"):
            df_review = pd.read_csv("data/output/incident_index.csv")
            valid_cats = [c.lower().replace(' ', '_') for c in selected_cats]
            display_df = df_review[df_review['category'].str.lower().str.replace(' ', '_').isin(valid_cats)]
            for cat in display_df['category'].unique():
                with st.expander(f"📁 {cat.upper()}", expanded=False):
                    st.data_editor(display_df[display_df['category'] == cat][['exhibit_id', 'dt', 'evidence_quote', 'reasoning']], use_container_width=True)
            if st.button("📋 GENERATE FINAL WORD REPORT"):
                export_to_word("data/output/incident_index.csv", "data/output/Custody_Report.docx")
                st.balloons()

    if os.path.exists("data/output/Custody_Report.docx"):
        with open("data/output/Custody_Report.docx", "rb") as file:
            st.sidebar.download_button("📥 DOWNLOAD COURT REPORT", data=file, file_name="Certified_Custody_Exhibits.docx")
else:
    st.info("👋 Welcome. Please upload your text export (CSV) in the sidebar to begin.")