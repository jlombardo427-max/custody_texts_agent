import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def export_to_word(incident_csv, output_path):
    """
    STRICT v3.8.4 Litigation-Ready Word Export:
    - PASS 1: Summary Exhibit Index (TOC) for quick judicial review.
    - PASS 2: Detailed Evidence Cards grouped by Legal Category.
    - N.J.R.E. 901 Traceability: Injects Row # and Message ID.
    """
    if not os.path.exists(incident_csv):
        print(f"Error: {incident_csv} not found.")
        return False

    # Load and prepare data
    df = pd.read_csv(incident_csv)
    df['dt'] = pd.to_datetime(df['dt'])
    df = df.sort_values(by=['category', 'dt'])

    doc = Document()

    # --- 1. LEGAL HEADER & CERTIFICATION ---
    title = doc.add_heading('CERTIFIED REPORT OF CUSTODY INCIDENTS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(f"Generated on: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "This report summarizes identified incidents from digital communication logs "
        "cross-referenced against NJSA 9:2-4 legal standards. Each exhibit is indexed "
        "to source data for forensic verification."
    )

    # --- 2. PASS 1: EXHIBIT INDEX SUMMARY (TOC) ---
    doc.add_heading('Exhibit Index Summary', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text = 'Exhibit ID', 'Date', 'Legal Category'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            paragraph.runs[0].bold = True

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['exhibit_id'])
        row_cells[1].text = row['dt'].strftime('%Y-%m-%d %H:%M')
        row_cells[2].text = str(row['category']).upper().replace('_', ' ')

    # --- 3. PASS 2: DETAILED EVIDENCE CHAPTERS ---
    for category in df['category'].unique():
        doc.add_page_break()
        chapter_title = category.replace('_', ' ').upper()
        doc.add_heading(f"CHAPTER: {chapter_title}", level=1)
        
        cat_df = df[df['category'] == category]
        
        for _, row in cat_df.iterrows():
            # Exhibit ID Header
            p = doc.add_paragraph()
            exhibit_run = p.add_run(f"EXHIBIT {str(row['exhibit_id']).upper()}")
            exhibit_run.bold = True
            exhibit_run.font.size = Pt(13)

            # Authenticity Footer (N.J.R.E. 901 Paper Trail)
            metadata = doc.add_paragraph()
            meta_text = (f"VERIFIED SOURCE DATA | Row: {row.get('raw_row_number', 'N/A')} | "
                         f"Date: {row['dt'].strftime('%Y-%m-%d %H:%M')} | "
                         f"ID: {row.get('message_id', 'N/A')}")
            meta_run = metadata.add_run(meta_text)
            meta_run.font.size = Pt(8)
            meta_run.italic = True
            metadata.paragraph_format.left_indent = Inches(0.5)

            # Evidence Text (Blockquote style)
            doc.add_paragraph("Evidence Quote:").runs[0].bold = True
            quote = doc.add_paragraph()
            quote.add_run(f"\"{row.get('evidence_quote', row.get('text', ''))}\"").italic = True
            quote.paragraph_format.left_indent = Inches(0.75)
            
            # AI Reasoning
            reason = doc.add_paragraph()
            reason_run = reason.add_run(f"Legal Reasoning: {row['reasoning']}")
            reason_run.bold = True
            reason.paragraph_format.left_indent = Inches(0.5)
            
            doc.add_paragraph("_" * 60) # Visual separator

    # --- 4. FOOTER ---
    section = doc.sections[0]
    footer = section.footer
    footer.paragraphs[0].text = "STRICT v3.8.4 | Giuseppe Lombardo - Confidential Legal Work Product | Page "

    # Save logic
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Final Report Generated: {output_path}")
    return True

if __name__ == "__main__":
    export_to_word("data/output/incident_index.csv", "data/output/Certified_Custody_Report.docx")